import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
import os
import imaplib
import email
from email.header import decode_header
from datetime import datetime

# CONFIGURAR GEMINI API
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
modelo = genai.GenerativeModel("gemini-2.0-flash")

st.set_page_config(page_title="Recrutamento IA com Gemini", layout="wide")
st.title("Análise Inteligente de Currículos com AI")

vaga = st.text_area("📌 Descreva a vaga:", height=150)
uploaded_files = st.file_uploader("📎 Envie os currículos (PDF, DOCX ou TXT):", type=["pdf", "docx", "txt"], accept_multiple_files=True)

# Pasta de base local
BASE_DIR = "curriculos_base"
os.makedirs(BASE_DIR, exist_ok=True)

# Função para baixar currículos do e-mail
def baixar_curriculos_email(email_user, email_pass):
    st.info("🔄 Conectando ao e-mail...")
    mail = imaplib.IMAP4_SSL("imap.gmail.com")
    mail.login(email_user, email_pass)
    mail.select("inbox")

    # Filtra e-mails com possíveis currículos (assunto ou anexos típicos)
    status, messages = mail.search(None, '(UNSEEN SUBJECT "curriculo" SUBJECT "cv" SUBJECT "vaga" SUBJECT "trabalho")')
    messages = messages[0].split()

    for num in messages:
        res, data = mail.fetch(num, '(RFC822)')
        raw_email = data[0][1]
        msg = email.message_from_bytes(raw_email)

        for part in msg.walk():
            if part.get_content_maintype() == 'multipart':
                continue
            if part.get('Content-Disposition') is None:
                continue

            filename = part.get_filename()
            if filename and filename.lower().endswith((".pdf", ".docx", ".txt")):
                filepath = os.path.join(BASE_DIR, filename)
                if not os.path.exists(filepath):
                    with open(filepath, 'wb') as f:
                        f.write(part.get_payload(decode=True))

    mail.logout()
    st.success(f"✅ Currículos atualizados em '{BASE_DIR}'")

# UI de atualização por e-mail
with st.expander("📥 Atualizar base de currículos por e-mail"):
    email_user = st.text_input("E-mail", placeholder="seu.email@gmail.com")
    email_pass = st.text_input("Senha (ou senha de app)", type="password")
    if st.button("🔄 Buscar currículos do e-mail"):
        if email_user and email_pass:
            baixar_curriculos_email(email_user, email_pass)
        else:
            st.warning("Preencha o e-mail e a senha!")

# Processar currículos da base local + uploads
curriculos_lidos = []

# Ler da base local
for filename in os.listdir(BASE_DIR):
    path = os.path.join(BASE_DIR, filename)
    if filename.lower().endswith((".pdf", ".docx", ".txt")):
        with open(path, 'rb') as file:
            texto = ""
            if filename.endswith(".pdf"):
                try:
                    reader = PdfReader(file)
                    texto = "\n".join([page.extract_text() or "" for page in reader.pages])
                except:
                    texto = ""
            elif filename.endswith(".docx"):
                doc = Document(file)
                texto = "\n".join([p.text for p in doc.paragraphs])
            elif filename.endswith(".txt"):
                texto = file.read().decode("utf-8")

            if texto.strip():
                curriculos_lidos.append((filename, texto))

# Ler dos uploads manuais
if uploaded_files:
    for arquivo in uploaded_files:
        texto = ""
        if arquivo.type == "application/pdf":
            try:
                reader = PdfReader(arquivo)
                texto = "\n".join([page.extract_text() or "" for page in reader.pages])
            except:
                texto = ""
        elif arquivo.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(arquivo)
            texto = "\n".join([p.text for p in doc.paragraphs])
        elif arquivo.type == "text/plain":
            texto = arquivo.read().decode("utf-8")

        if texto.strip():
            curriculos_lidos.append((arquivo.name, texto))

# Funções auxiliares

def extrair_palavras_chave(texto):
    return set(re.findall(r'\b\w{4,}\b', texto.lower()))

def filtrar_por_palavras_chave(curriculos, vaga_texto):
    palavras_vaga = extrair_palavras_chave(vaga_texto)
    return [(nome, texto) for nome, texto in curriculos if palavras_vaga & extrair_palavras_chave(texto)]

def ranquear_por_tfidf(curriculos, vaga_texto, top_n=5):
    documentos = [vaga_texto] + [texto for _, texto in curriculos]
    nomes = [None] + [nome for nome, _ in curriculos]
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documentos)
    similaridades = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
    indices_top = similaridades.argsort()[::-1][:top_n]
    return [(nomes[i+1], curriculos[i][1], similaridades[i]) for i in indices_top]

resultados = []

if st.button("🚀 Analisar Currículos") and vaga and curriculos_lidos:
    with st.spinner("Analisando..."):
        # Pré-filtragem por palavras-chave
        filtrados = filtrar_por_palavras_chave(curriculos_lidos, vaga)
        # Ranqueamento com TF-IDF (top 5)
        selecionados = ranquear_por_tfidf(filtrados, vaga, top_n=5)

        for nome, texto, score in selecionados:
            prompt = f"""
Compare o seguinte currículo com a vaga abaixo.

Vaga:
{vaga}

Currículo:
{texto}

Responda com:
1. Uma pontuação de 0 a 100 sobre compatibilidade.
2. Um resumo do porquê o perfil combina ou não com a vaga.
"""
            response = modelo.generate_content(prompt)
            resultado = response.text
            resultados.append((nome, resultado))

            st.markdown(f"### 📄 {nome} — Similaridade TF-IDF: `{score:.2f}`")
            st.write(resultado)
            st.divider()

    # Criar PDF com resultados
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawCentredString(width / 2, height - 2 * cm, "Relatório de Análise de Currículos")
    y = height - 3 * cm

    from reportlab.lib.utils import simpleSplit

    for nome, texto in resultados:
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(2 * cm, y, f"Currículo: {nome}")
        y -= 0.6 * cm
        pdf.setFont("Helvetica", 11)
        for linha in texto.split('\n'):
            partes = simpleSplit(linha, 'Helvetica', 11, width - 4 * cm)
            for parte in partes:
                if y < 2 * cm:
                    pdf.showPage()
                    y = height - 2 * cm
                    pdf.setFont("Helvetica", 11)
                pdf.drawString(2 * cm, y, parte)
                y -= 0.5 * cm
        y -= 1 * cm

    pdf.save()
    buffer.seek(0)

    st.download_button("📥 Baixar Relatório em PDF", buffer, file_name="relatorio_analise_curriculos.pdf", mime="application/pdf")
